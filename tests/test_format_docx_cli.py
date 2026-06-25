from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
FORMAT_CLI = PROJECT_ROOT / "scripts" / "format_docx.py"
PUBLIC_SERVICE_FIXTURE = PROJECT_ROOT / "tests" / "fixtures" / "docx" / "public_service_guide_single_paragraph.docx"


def run_format_cli(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(FORMAT_CLI), *args],
        cwd=PROJECT_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )


def make_docx(path: Path, lines: list[str]) -> Path:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(str(path))
    return path


def footer_xml(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(section.footer._element.xml for section in doc.sections)


def footer_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for section in doc.sections for paragraph in section.footer.paragraphs)


def document_xml(path: Path) -> str:
    doc = Document(str(path))
    return doc._element.xml


def document_body_block_texts(path: Path) -> list[str]:
    doc = Document(str(path))
    body = doc.element.body
    paragraph_index = 0
    table_index = 0
    blocks: list[str] = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = doc.paragraphs[paragraph_index].text.strip()
            paragraph_index += 1
            if text:
                blocks.append(text)
        elif child.tag == qn("w:tbl"):
            table = doc.tables[table_index]
            table_index += 1
            first_row = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
            blocks.append(f"TABLE: {first_row}")
    return blocks


def table_grid_widths(path: Path) -> list[int]:
    doc = Document(str(path))
    grid = doc.tables[0]._tbl.tblGrid
    return [int(column.get(qn("w:w"))) for column in grid.gridCol_lst]


def test_confident_notice_fixture_generates_docx_and_report(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "notice_source.docx",
        [
            "关于开展安全生产检查的通知",
            "各部门：",
            "一、总体要求",
            "请按要求开展安全生产检查。",
            "（一）检查范围",
            "各部门自查。",
            "测试单位",
            "2026年6月19日",
        ],
    )
    output_path = tmp_path / "notice.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    assert report_path.exists()
    assert "saved=" in result.stdout
    assert "report=" in result.stdout

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["input"] == input_path.name
    assert report["output"] == output_path.name
    assert report["profile_id"] == "standard-party-government"
    assert report["doc_type"] == "通知"
    assert report["run_config"] == {
        "profile_id": "standard-party-government",
        "doc_type": "通知",
        "text_normalization": "keep_en_boundary",
        "space_mode": "keep_en_boundary",
        "include_local_paths": False,
    }
    assert report["structure"]["title_indices"]
    assert report["operations"]


def test_report_can_include_local_paths_when_explicitly_requested(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "notice_source.docx",
        [
            "关于开展安全生产检查的通知",
            "各部门：",
            "请按要求开展安全生产检查。",
        ],
    )
    output_path = tmp_path / "notice.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--include-local-paths")

    assert result.returncode == 0, result.stderr
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["input"] == str(input_path)
    assert report["output"] == str(output_path)
    assert report["run_config"]["include_local_paths"] is True


def test_ambiguous_fixture_returns_2_and_does_not_generate_docx(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "ambiguous_source.docx",
        [
            "关于数据治理有关事项的材料",
            "现将有关情况报告如下。",
            "请贵单位协助支持。",
        ],
    )
    output_path = tmp_path / "ambiguous.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 2
    assert not output_path.exists()
    assert not report_path.exists()
    assert "你希望怎么处理" in result.stdout
    assert "通用正式文本" in result.stdout


def test_page_numbers_flag_inserts_word_page_field_and_reports_action(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "notice_source.docx",
        [
            "关于开展安全生产检查的通知",
            "各部门：",
            "请按要求开展安全生产检查。",
            "某某办公室",
            "2026年6月21日",
        ],
    )
    output_path = tmp_path / "notice_with_page_number.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--page-numbers")

    assert result.returncode == 0, result.stderr
    xml = footer_xml(output_path)
    assert "PAGE" in xml
    assert "fldChar" in xml

    report = json.loads(report_path.read_text(encoding="utf-8"))
    page_ops = [operation for operation in report["operations"] if operation["kind"] == "page_number"]
    assert page_ops
    assert page_ops[0]["params"]["existing_page_number"] is False
    assert page_ops[0]["params"]["action"] == "inserted"


def test_page_numbers_flag_preserves_non_page_footer_and_skips_insertion(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_with_footer.docx"
    source = Document()
    source.add_paragraph("关于开展安全检查的通知")
    source.add_paragraph("各部门：")
    source.add_paragraph("请按要求开展安全检查。")
    source.add_paragraph("某某办公室")
    source.add_paragraph("2026年6月21日")
    source.sections[0].footer.paragraphs[0].text = "内部资料"
    source.save(str(input_path))

    output_path = tmp_path / "notice_with_footer_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--page-numbers")

    assert result.returncode == 0, result.stderr
    assert "内部资料" in footer_text(output_path)
    assert "PAGE" not in footer_xml(output_path)

    report = json.loads(report_path.read_text(encoding="utf-8"))
    page_ops = [operation for operation in report["operations"] if operation["kind"] == "page_number"]
    assert page_ops
    assert page_ops[0]["params"]["existing_non_page_footer"] is True
    assert page_ops[0]["params"]["action"] == "skipped_non_page_footer"


def test_format_tables_flag_preserves_table_and_formats_cell_text(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_with_table.docx"
    source = Document()
    source.add_paragraph("关于报送整改台账的通知")
    source.add_paragraph("各部门：")
    source.add_paragraph("请按表格要求报送整改台账。")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "事项"
    table.cell(0, 1).text = "责任部门"
    table.cell(1, 0).text = "安全检查"
    table.cell(1, 1).text = "综合部"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = None
    source.save(str(input_path))

    output_path = tmp_path / "notice_with_table_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--format-tables")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    assert len(output.tables) == 1
    first_run = output.tables[0].cell(0, 0).paragraphs[0].runs[0]
    assert first_run.text == "事项"
    assert first_run.font.name == "仿宋_GB2312"
    assert first_run.font.size.pt == 16

    report = json.loads(report_path.read_text(encoding="utf-8"))
    table_ops = [operation for operation in report["operations"] if operation["kind"] == "table_format"]
    assert table_ops
    assert table_ops[0]["params"]["table_count"] == 1
    assert table_ops[0]["params"]["action"] == "formatted_inline"


def test_formatting_does_not_treat_numbered_colon_heading_as_recipient_or_drop_tables(tmp_path: Path) -> None:
    input_path = tmp_path / "public_service_rule_source.docx"
    source = Document()
    for text in [
        "公共服务事项办理指引",
        "一、公共服务事项定义与识别",
        "1.事项定义：公共服务事项是指面向社会公众提供的咨询、受理、审查、反馈等服务事项。",
        "2.识别与公开标准：服务机构应通过统一服务平台，定期梳理高频事项，将其列入办事指南清单进行专项管理。",
        "二、组织职责与权限分工",
        "1.办事指南编制部门：作为归口管理部门，负责制定事项清单，建立信息平台，并对办理情况进行定期评估。",
        "2.服务窗口：负责现场咨询、材料接收、一次告知和办理进度反馈。",
        "3.技术支持部门：负责平台运行、账号权限、电子材料上传和系统故障处理。",
        "4.业务办理部门：负责事项审核、结果确认、资料归档及问题反馈。",
        "三、信息发布规范",
        "1.信息发布规范：事项信息须包含事项名称、申请条件、办理材料、办理地点、办理机构和办理流程等关键要素。",
        "2.动态更新机制：当政策依据、材料要求或办理地址发生变化时，发布单位应在24小时内完成信息同步。",
        "四、“应办未办”理由审核标准",
        "1.主动提醒规则：申请人在发起办理申请前，系统将自动匹配相关办事指南。若存在同类事项，原则上应优先引导线上办理。",
        "2.合理退回理由：",
        "（一）材料差异处理逻辑",
    ]:
        source.add_paragraph(text)
    table = source.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "交易主体场景"
    table.cell(0, 1).text = "价值门槛"
    table.cell(0, 2).text = "定价基础"
    table.cell(0, 3).text = "评估要求"
    table.cell(1, 0).text = "跨单位"
    table.cell(1, 1).text = "一般物资"
    table.cell(1, 2).text = "评估价"
    table.cell(1, 3).text = "按规则执行"
    source.add_paragraph("（二）核心概念定义与流程")
    source.save(str(input_path))
    output_path = tmp_path / "public_service_rule.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--assume-detected-type")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert texts[:15] == [
        "公共服务事项办理指引",
        "一、公共服务事项定义与识别",
        "1.事项定义：公共服务事项是指面向社会公众提供的咨询、受理、审查、反馈等服务事项。",
        "2.识别与公开标准：服务机构应通过统一服务平台，定期梳理高频事项，将其列入办事指南清单进行专项管理。",
        "二、组织职责与权限分工",
        "1.办事指南编制部门：作为归口管理部门，负责制定事项清单，建立信息平台，并对办理情况进行定期评估。",
        "2.服务窗口：负责现场咨询、材料接收、一次告知和办理进度反馈。",
        "3.技术支持部门：负责平台运行、账号权限、电子材料上传和系统故障处理。",
        "4.业务办理部门：负责事项审核、结果确认、资料归档及问题反馈。",
        "三、信息发布规范",
        "1.信息发布规范：事项信息须包含事项名称、申请条件、办理材料、办理地点、办理机构和办理流程等关键要素。",
        "2.动态更新机制：当政策依据、材料要求或办理地址发生变化时，发布单位应在24小时内完成信息同步。",
        "四、“应办未办”理由审核标准",
        "1.主动提醒规则：申请人在发起办理申请前，系统将自动匹配相关办事指南。若存在同类事项，原则上应优先引导线上办理。",
        "2.合理退回理由：",
    ]
    assert len(output.tables) == 1
    blocks = document_body_block_texts(output_path)
    assert blocks.index("（一）材料差异处理逻辑") < blocks.index("TABLE: 交易主体场景 | 价值门槛 | 定价基础 | 评估要求")
    assert blocks.index("TABLE: 交易主体场景 | 价值门槛 | 定价基础 | 评估要求") < blocks.index("（二）核心概念定义与流程")


def test_formatting_preserves_multiline_title_style(tmp_path: Path) -> None:
    input_path = make_docx(
            tmp_path / "multiline_title_source.docx",
            [
            "公共服务流程优化后的办理路径与",
            "原流程对比说明",
            "一、总体情况",
            "本材料说明流程差异。",
        ],
    )
    output_path = tmp_path / "multiline_title.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    assert non_empty[0].text.strip() == "公共服务流程优化后的办理路径与"
    assert non_empty[1].text.strip() == "原流程对比说明"
    for paragraph in non_empty[:2]:
        assert paragraph.alignment == 1
        assert paragraph.runs[0].font.name == "方正小标宋简体"
        assert paragraph.runs[0].font.size.pt == 22


def test_multiline_title_detection_does_not_swallow_front_matter_or_tables(tmp_path: Path) -> None:
    input_path = tmp_path / "prd_source.docx"
    source = Document()
    source.add_paragraph("供应链新增应急借料业务")
    source.add_paragraph("产品需求规格说明书")
    source.add_paragraph("版本记录")
    for index in range(4):
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = f"字段{index + 1}"
        table.cell(0, 1).text = "说明"
    source.save(str(input_path))
    output_path = tmp_path / "prd.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    assert non_empty[0].text.strip() == "供应链新增应急借料业务"
    assert non_empty[1].text.strip() == "产品需求规格说明书"
    assert non_empty[2].text.strip() == "版本记录"
    assert len(output.tables) == 4


def test_standard_text_document_uses_standard_branch_and_specific_layout(tmp_path: Path) -> None:
    input_path = tmp_path / "standard_text_source.docx"
    source = Document()
    for text in [
        "中华人民共和国",
        "电力企业团体标准配套稿",
        "供应链服务标准",
        "（征求意见稿）",
        "目次",
        "前    言",
        "1  范围",
        "3.1  服务事项",
        "前    言",
        "本标准依据相关规则起草。",
        "1  范围",
        "本文件规定了供应链服务范围。",
        "3.1  服务事项",
        "面向公众提供咨询、受理、审查、反馈等服务的事项。",
    ]:
        source.add_paragraph(text)
    table = source.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "项目"
    table.cell(0, 2).text = "要求"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "信息发布"
    table.cell(1, 2).text = "内容完整"
    source.save(str(input_path))
    output_path = tmp_path / "standard_text.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in non_empty]
    assert texts[0] == "中华人民共和国电力企业团体标准配套稿"
    body_preface_index = max(index for index, text in enumerate(texts) if text == "前    言")
    assert texts[body_preface_index + 1].startswith("本标准依据")

    toc_chapter = next(paragraph for paragraph in non_empty if paragraph.text.strip() == "1  范围")
    toc_clause = next(paragraph for paragraph in non_empty if paragraph.text.strip() == "3.1  服务事项")
    assert toc_chapter.runs[0].bold is True
    assert toc_clause.runs[0].bold is False
    assert table_grid_widths(output_path) == [1061, 3892, 3892]

    report = json.loads(report_path.read_text(encoding="utf-8"))
    standard_ops = [operation for operation in report["operations"] if operation["kind"] == "standard_text_format"]
    assert standard_ops
    assert report["doc_type"] == "标准规范文本"
    assert standard_ops[0]["params"]["merged_cover_label"] is True


def test_generic_formal_text_doc_type_formats_ambiguous_material_without_prompt(tmp_path: Path) -> None:
    input_path = tmp_path / "generic_material.docx"
    source = Document()
    for text in [
        "专项工作交流材料",
        "相关单位：",
        "一、基本情况",
        "有关工作正在稳步推进。",
        "（一）推进情况",
        "各项任务按照计划开展。",
        "1. 重点事项",
        "重点事项已形成阶段性成果。",
        "某某部门",
        "2026年6月23日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "generic_material_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--doc-type", "通用正式文本")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    assert "你希望" not in result.stdout

    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert texts[0] == "专项工作交流材料"
    assert texts[1] == "相关单位："
    assert texts[-2:] == ["某某部门", "2026年6月23日"]
    assert output.paragraphs[0].alignment == 1

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "通用正式文本"


def test_generic_formal_text_flag_formats_without_prompt(tmp_path: Path) -> None:
    input_path = tmp_path / "generic_material_flag.docx"
    source = Document()
    source.add_paragraph("专项工作交流材料")
    source.add_paragraph("一、基本情况")
    source.add_paragraph("有关工作正在推进。")
    source.save(str(input_path))

    output_path = tmp_path / "generic_material_flag_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    assert "你希望" not in result.stdout
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "通用正式文本"


def test_text_file_mode_writes_report_for_regular_multiline_text(tmp_path: Path) -> None:
    input_path = tmp_path / "generic_material.txt"
    input_path.write_text(
        "\n".join([
            "专项工作交流材料",
            "一、基本情况",
            "有关工作正在推进。",
        ]),
        encoding="utf-8",
    )
    output_path = tmp_path / "generic_material.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli("--text-file", str(input_path), "-o", str(output_path), "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    assert report_path.exists()
    assert "report=" in result.stdout
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["input"] == input_path.name
    assert report["output"] == output_path.name
    assert report["input_type"] == "text_file"
    assert report["doc_type"] == "通用正式文本"
    assert report["run_config"] == {
        "profile_id": "standard-party-government",
        "doc_type": "通用正式文本",
        "text_normalization": "keep_en_boundary",
        "space_mode": "keep_en_boundary",
        "include_local_paths": False,
    }
    assert report["glued_text_detected"] is False
    assert report["raw_line_count"] == 3
    assert report["output_paragraph_count"] == 3
    assert report["warnings"] == []


def test_text_file_mode_reports_glued_long_text_without_rebuilding_structure(tmp_path: Path) -> None:
    input_path = tmp_path / "center_warehouse.txt"
    input_path.write_text(
        "中心仓建设和管理工作指引"
        "为进一步落实《开展集团公司2025年区域中心仓建设工作》，规范中心仓建设与作业管理。"
        "基本原则目标导向：以提升物资集采效率、降低供应链成本为核心。"
        "设备大类 | 设备小类 | 相关设备 | 功能定位"
        "安防与监控 | 网络硬件 | 温湿度监控仪、POE交换机 | 基础环境监控与安全防护"
        + "中心仓作业管理。" * 140,
        encoding="utf-8",
    )
    output_path = tmp_path / "center_warehouse.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(
        "--text-file",
        str(input_path),
        "-o",
        str(output_path),
        "--report",
        "--generic-formal-text",
        "--title",
        "中心仓建设和管理工作指引",
        "--format-tables",
    )

    assert result.returncode == 0, result.stderr
    assert "warning=glued_plain_text_detected" in result.stdout
    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert len(texts) == 2
    assert len(output.tables) == 0

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["input_type"] == "text_file"
    assert report["glued_text_detected"] is True
    assert report["raw_line_count"] == 1
    assert report["output_paragraph_count"] == 2
    assert report["warnings"] == [
        "glued_plain_text_detected: text-file input has too few line breaks; structure and tables were not recovered"
    ]
    assert any(operation["kind"] == "text_file_format" for operation in report["operations"])


def test_report_records_requested_space_mode(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "notice_source.docx",
        [
            "关于开展安全生产检查的通知",
            "各部门：",
            "请按要求开展安全生产检查。",
        ],
    )
    output_path = tmp_path / "notice_remove_spaces.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--space-mode", "remove_all")

    assert result.returncode == 0, result.stderr
    assert "text_normalization=remove_all" in result.stdout
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["run_config"]["space_mode"] == "remove_all"
    text_ops = [operation for operation in report["operations"] if operation["kind"] == "text_normalization"]
    assert text_ops
    assert {operation["params"]["space_mode"] for operation in text_ops} == {"remove_all"}


def test_glued_single_paragraph_report_is_split_into_reasonable_blocks(tmp_path: Path) -> None:
    input_path = tmp_path / "glued_report_source.docx"
    source = Document()
    paragraph = source.add_paragraph(
        "公共服务事项办理规范自查报告"
        "为提高公共服务事项信息的规范性和专业性，进一步提升事项办理管理水平，组织开展了自查工作。"
        "现将自查情况报告如下。"
        "存在的问题"
        "经全面排查，部分事项的办理材料字段存在表述不统一问题，影响事项信息的统一展示。"
        "解决措施"
        "针对材料表述问题，计划采取分步整改的方式，将不统一表述调整为规范表述。"
        "时间计划"
        "按期完成材料字段批量核对工作，并完成相关整改复核验收。"
        "服务专班"
    )
    paragraph.alignment = 1
    source.save(str(input_path))

    output_path = tmp_path / "glued_report_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert texts == [
        "公共服务事项办理规范自查报告",
        "为提高公共服务事项信息的规范性和专业性，进一步提升事项办理管理水平，组织开展了自查工作。现将自查情况报告如下。",
        "存在的问题",
        "经全面排查，部分事项的办理材料字段存在表述不统一问题，影响事项信息的统一展示。",
        "解决措施",
        "针对材料表述问题，计划采取分步整改的方式，将不统一表述调整为规范表述。",
        "时间计划",
        "按期完成材料字段批量核对工作，并完成相关整改复核验收。",
        "服务专班",
    ]
    assert output.paragraphs[0].alignment == 1
    assert output.paragraphs[1].alignment != 1
    for heading_text in ["存在的问题", "解决措施", "时间计划"]:
        heading = next(paragraph for paragraph in output.paragraphs if paragraph.text.strip() == heading_text)
        assert heading.paragraph_format.first_line_indent.pt == 32
        assert heading.runs[0].font.name == "黑体"
        assert heading.runs[0].font.bold is True

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "报告"
    assert report["structure"]["title_indices"] == [0]


def test_unnumbered_headings_and_subtitle_are_structured(tmp_path: Path) -> None:
    input_path = tmp_path / "unnumbered_headings_source.docx"
    input_path = make_docx(
        input_path,
        [
            "公共服务事项办理规范自查报告",
            "整改情况说明",
            "为提高事项办理管理水平，组织开展了自查工作。",
            "存在的问题",
            "部分事项缺少示范文本，影响申请人准确理解办理材料要求。",
            "解决措施",
            "已安排专人对接示范文本的核查工作，并反馈至对应部门补充说明。",
            "时间计划",
            "按期完成示范文本清单排查、部门对接及材料补充更新。",
            "服务专班",
        ],
    )
    output_path = tmp_path / "unnumbered_headings_formatted.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    assert [paragraph.text.strip() for paragraph in non_empty[:2]] == ["公共服务事项办理规范自查报告", "整改情况说明"]
    assert non_empty[0].alignment == 1
    assert non_empty[1].alignment == 1
    assert non_empty[1].runs[0].font.name == "方正小标宋简体"
    for heading_text in ["存在的问题", "解决措施", "时间计划"]:
        heading = next(paragraph for paragraph in non_empty if paragraph.text.strip() == heading_text)
        assert heading.paragraph_format.first_line_indent.pt == 32
        assert heading.runs[0].font.name == "黑体"
        assert heading.runs[0].font.bold is True
    signature = non_empty[-1]
    assert signature.text.strip() == "服务专班"
    assert signature.runs[0].font.name == "仿宋_GB2312"
    assert signature.runs[0].font.bold is False


def test_spaced_single_paragraph_report_recovers_nested_hierarchy(tmp_path: Path) -> None:
    input_path = tmp_path / "spaced_glued_report_source.docx"
    source = Document()
    source.add_paragraph(
        "公共服务事项办理规范自查报告 "
        "为提高公共服务事项信息的规范性和专业性，进一步提升事项办理管理水平，根据事项公开相关要求，组织开展了公共服务事项办理规范自查工作。现将自查情况报告如下。 "
        "存在的问题 "
        "材料清单不符合事项公开规范 "
        "经全面排查，部分事项的材料清单字段存在表述不统一问题，影响事项信息标准化管理及统一展示。 "
        "办理时限存在异常 "
        "在自查过程中发现，部分事项的承诺办理时限与实际流程要求差异较大。 "
        "部分事项缺少示范文本 "
        "示范文本是申请人对办理材料要求进行直观判断的重要依据。 "
        "解决措施 "
        "材料清单不规范问题整改方案 "
        "针对材料清单表述问题，计划采取分步整改的方式。 "
        "办理时限异常问题整改方案 "
        "针对已标记的疑似异常记录，事项管理员将逐一联系对应部门进行时限核实。 "
        "示范文本缺失问题整改方案 "
        "已安排专人对接缺少示范文本事项的核查工作。 "
        "时间计划 "
        "按期完成材料清单批量核对申请的提交工作。 "
        "后续将定期跟踪整改进度，确保各项整改措施按计划推进落实。 "
        "服务专班"
    )
    source.save(str(input_path))
    output_path = tmp_path / "spaced_glued_report_formatted.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in non_empty]
    assert texts == [
        "公共服务事项办理规范自查报告",
        "为提高公共服务事项信息的规范性和专业性，进一步提升事项办理管理水平，根据事项公开相关要求，组织开展了公共服务事项办理规范自查工作。现将自查情况报告如下。",
        "一、存在的问题",
        "（一）材料清单不符合事项公开规范",
        "经全面排查，部分事项的材料清单字段存在表述不统一问题，影响事项信息标准化管理及统一展示。",
        "（二）办理时限存在异常",
        "在自查过程中发现，部分事项的承诺办理时限与实际流程要求差异较大。",
        "（三）部分事项缺少示范文本",
        "示范文本是申请人对办理材料要求进行直观判断的重要依据。",
        "二、解决措施",
        "（一）材料清单不规范问题整改方案",
        "针对材料清单表述问题，计划采取分步整改的方式。",
        "（二）办理时限异常问题整改方案",
        "针对已标记的疑似异常记录，事项管理员将逐一联系对应部门进行时限核实。",
        "（三）示范文本缺失问题整改方案",
        "已安排专人对接缺少示范文本事项的核查工作。",
        "三、时间计划",
        "按期完成材料清单批量核对申请的提交工作。",
        "后续将定期跟踪整改进度，确保各项整改措施按计划推进落实。",
        "服务专班",
    ]
    level1 = [paragraph for paragraph in non_empty if paragraph.text.strip().startswith(("一、", "二、", "三、"))]
    level2 = [paragraph for paragraph in non_empty if paragraph.text.strip().startswith(("（一）", "（二）", "（三）"))]
    assert {paragraph.runs[0].font.name for paragraph in level1} == {"黑体"}
    assert {paragraph.paragraph_format.first_line_indent.pt for paragraph in level1} == {32}
    assert {paragraph.runs[0].font.name for paragraph in level2} == {"楷体_GB2312"}
    assert non_empty[-1].runs[0].font.name == "仿宋_GB2312"


def test_public_service_single_paragraph_uses_scripted_chapter_recovery(tmp_path: Path) -> None:
    input_path = tmp_path / "public_service_source.docx"
    source = Document()
    source.add_paragraph(
        "公共服务事项办理指引"
        "公共服务事项是指面向社会公众提供的咨询、受理、审查、反馈等服务事项，办理过程应当公开透明、便捷高效。"
        "线上服务平台应集中展示事项名称、设定依据、申请条件、办理材料、办理地点、办理机构、收费标准、办理时间、联系电话和办理流程。"
        "办事指南编制部门负责事项内容维护、流程说明、示范文本发布和常见问题更新。"
        "服务窗口负责现场咨询、材料接收、一次告知和办理进度反馈。"
        "技术支持部门负责平台运行、账号权限、电子材料上传和系统故障处理。"
        "办理材料清单应列明材料名称、来源渠道、纸质或电子形式、份数要求、签名签章要求和示范样例。"
        "申请人在提交申请前，应先核对事项条件和材料清单。"
        "受理标准主要依据事项公开条件、材料完整性和申请主体资格确定。"
        "补正告知是指受理人员发现材料不完整或者格式不符合要求时，一次性告知需要补正的内容。"
        "审查办理是通过材料核验、业务复核和必要的现场核查，确认申请事项是否符合办理条件。"
        "结果送达是指通过现场领取、邮寄送达或者线上下载等方式向申请人反馈办理结果。"
        "为提高办理效率，服务机构应对高频事项设置快速办理通道。"
        "所有办理过程应在统一服务平台记录申请、受理、补正、审查、办结和评价等关键节点。"
        "监督评价方面，申请人可以对办理时限、服务态度、结果反馈和问题处理情况进行评价。"
        "资料归档适用“谁办理、谁归档”的原则。"
        "改进机制层面，服务机构应定期汇总咨询问题、退回原因和用户评价，持续优化办事指南和办理流程。"
    )
    source.save(str(input_path))
    output_path = tmp_path / "public_service_formatted.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--assume-detected-type")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in non_empty]
    expected_chapters = [
        "一、职责分工",
        "二、办理材料清单",
        "三、申请前核对",
        "四、受理标准",
        "五、补正告知",
        "六、审查办理",
        "七、结果送达",
        "八、高频事项快办",
        "九、全流程记录",
        "十、监督评价",
        "十一、资料归档",
        "十二、问题汇总",
        "十三、持续改进",
    ]
    assert texts[0] == "公共服务事项办理指引"
    assert all(chapter in texts for chapter in expected_chapters)
    assert not any(text in {"职责分工", "办理材料", "监督评价", "持续改进"} for text in texts)
    level1 = [paragraph for paragraph in non_empty if paragraph.text.strip() in expected_chapters]
    assert len(level1) == len(expected_chapters)
    assert {paragraph.runs[0].font.name for paragraph in level1} == {"黑体"}
    assert {paragraph.paragraph_format.first_line_indent.pt for paragraph in level1} == {32}


def test_public_service_real_single_paragraph_fixture_recovers_structure(tmp_path: Path) -> None:
    assert PUBLIC_SERVICE_FIXTURE.exists()

    entry_modes = [
        ("assume_detected", ["--assume-detected-type"]),
        ("generic_flag", ["--generic-formal-text"]),
        ("generic_doc_type", ["--doc-type", "通用正式文本"]),
    ]
    expected_fragments = [
        "公共服务事项办理指引",
        "受理标准",
        "监督评价",
        "持续改进",
    ]

    for name, extra_args in entry_modes:
        output_path = tmp_path / f"{name}_formatted.docx"
        report_path = output_path.with_suffix(".report.json")

        result = run_format_cli(str(PUBLIC_SERVICE_FIXTURE), "-o", str(output_path), "--report", *extra_args)

        assert result.returncode == 0, result.stderr
        output = Document(str(output_path))
        texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
        assert len(texts) > 10
        assert texts[0] == "公共服务事项办理指引"
        assert all(any(fragment in text for text in texts) for fragment in expected_fragments)
        assert not (len(texts) == 1 and len(texts[0]) > 1000)

        report = json.loads(report_path.read_text(encoding="utf-8"))
        recovery_ops = [operation for operation in report["operations"] if operation["kind"] == "chapter_recovery"]
        assert recovery_ops
        assert recovery_ops[0]["params"]["method"] == "public_service_guide"


def test_generic_formal_text_flag_overrides_standard_spec_auto_detection(tmp_path: Path) -> None:
    input_path = tmp_path / "standard_like_without_toc.docx"
    source = Document()
    for text in [
        "发电企业供应链服务标准",
        "前    言",
        "本文件说明相关服务要求。",
        "1  范围",
        "本文件规定了供应链服务范围。",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "standard_like_as_generic.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "通用正式文本"


def test_generate_toc_flag_adds_word_toc_field_for_clear_headings(tmp_path: Path) -> None:
    input_path = tmp_path / "report_with_headings.docx"
    source = Document()
    for text in [
        "关于安全生产整改情况的报告",
        "集团公司：",
        "一、总体情况",
        "本段介绍总体情况。",
        "（一）隐患排查情况",
        "本段介绍排查情况。",
        "1. 重点问题",
        "本段介绍重点问题。",
        "二、下一步工作",
        "本段介绍下一步工作。",
        "某某公司",
        "2026年6月21日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "report_with_toc.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generate-toc")

    assert result.returncode == 0, result.stderr
    xml = document_xml(output_path)
    assert "TOC" in xml
    assert "fldChar" in xml
    assert "outlineLvl" in xml

    report = json.loads(report_path.read_text(encoding="utf-8"))
    toc_ops = [operation for operation in report["operations"] if operation["kind"] == "toc_generation"]
    assert toc_ops
    assert toc_ops[0]["params"]["action"] == "generated"
    assert toc_ops[0]["params"]["heading_count"] >= 4


def test_generate_toc_flag_skips_when_headings_are_not_clear(tmp_path: Path) -> None:
    input_path = tmp_path / "short_notice.docx"
    source = Document()
    for text in [
        "关于开展安全检查的通知",
        "各部门：",
        "请按要求开展安全检查。",
        "某某办公室",
        "2026年6月21日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "short_notice_no_toc.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generate-toc")

    assert result.returncode == 0, result.stderr
    assert "TOC" not in document_xml(output_path)

    report = json.loads(report_path.read_text(encoding="utf-8"))
    toc_ops = [operation for operation in report["operations"] if operation["kind"] == "toc_generation"]
    assert toc_ops
    assert toc_ops[0]["params"]["action"] == "skipped_unclear_headings"
    assert toc_ops[0]["params"]["heading_count"] == 0


def test_format_imprint_flag_formats_existing_imprint_and_reports_action(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_with_imprint.docx"
    source = Document()
    for text in [
        "关于开展安全检查的通知",
        "各部门：",
        "请按要求开展安全检查。",
        "某某办公室",
        "2026年6月21日",
        "某某办公室                        2026年6月21日印发",
    ]:
        paragraph = source.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = None
    source.save(str(input_path))

    output_path = tmp_path / "notice_with_imprint_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--format-imprint")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    imprint = next(paragraph for paragraph in output.paragraphs if "印发" in paragraph.text)
    assert imprint.paragraph_format.first_line_indent.pt == 0
    assert imprint.runs[0].font.name == "仿宋_GB2312"
    assert imprint.runs[0].font.size.pt == 16

    report = json.loads(report_path.read_text(encoding="utf-8"))
    imprint_ops = [operation for operation in report["operations"] if operation["kind"] == "imprint_format"]
    assert imprint_ops
    assert imprint_ops[0]["params"]["action"] == "formatted"
    assert imprint_ops[0]["params"]["imprint_count"] == 1


def test_format_imprint_flag_reports_no_imprint_without_adding_one(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_without_imprint.docx"
    source = Document()
    for text in [
        "关于开展安全检查的通知",
        "各部门：",
        "请按要求开展安全检查。",
        "某某办公室",
        "2026年6月21日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "notice_without_imprint_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--format-imprint")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    assert not any("印发" in paragraph.text for paragraph in output.paragraphs)

    report = json.loads(report_path.read_text(encoding="utf-8"))
    imprint_ops = [operation for operation in report["operations"] if operation["kind"] == "imprint_format"]
    assert imprint_ops
    assert imprint_ops[0]["params"]["action"] == "no_imprint"
    assert imprint_ops[0]["params"]["imprint_count"] == 0
