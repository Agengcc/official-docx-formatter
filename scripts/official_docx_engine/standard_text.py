"""Formatting support for Chinese standard-specification style DOCX files."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Union

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class ParaBlock:
    text: str


@dataclass(frozen=True)
class TableBlock:
    rows: list[list[str]]


Block = Union[ParaBlock, TableBlock]


@dataclass(frozen=True)
class StandardTextResult:
    action: str
    toc_count: int
    table_count: int
    merged_cover_label: bool


BODY_FONT = "body"
HEADING_FONT = "level1"
SONG_FONT = "page_number"
TITLE_FONT = "title"

STANDARD_MARKER_RE = re.compile(r"^1\s+范围$")
CHAPTER_RE = re.compile(r"^\d+\s+.+")
CLAUSE_RE = re.compile(r"^\d+(?:\.\d+)+\s+.+")
LIST_ITEM_RE = re.compile(r"^[a-z]）")


def looks_like_standard_text(path: str | Path) -> bool:
    """Detect enterprise/national standard-specification structure."""

    document = Document(str(path))
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    compact_texts = {re.sub(r"\s+", "", text) for text in texts}
    return (
        bool({"目次", "前言"} & compact_texts)
        and any(STANDARD_MARKER_RE.match(" ".join(text.split())) for text in texts)
        and any("标准" in text for text in texts[:8])
    )


def build_standard_text_document(
    input_path: str | Path,
    profile: dict[str, Any],
) -> tuple[Document, StandardTextResult]:
    """Build a formatted standard-specification document while preserving source text order."""

    source = Document(str(input_path))
    cover, toc, body = _split_standard_document(list(_iter_blocks(source)))
    cover, merged_cover_label = _merge_cover_standard_label(cover)

    document = Document()
    _set_page(document.sections[0], profile)

    for index, text in enumerate(cover):
        _add_cover_text(document, index, text, profile)

    document.add_page_break()

    for index, text in enumerate(toc):
        _add_toc_text(document, index, text, profile)

    document.add_page_break()

    table_count = 0
    for block in body:
        if isinstance(block, TableBlock):
            table_count += 1
        _add_body_block(document, block, profile)

    for section in document.sections:
        _set_page(section, profile)
        _add_page_number(section, profile)

    return document, StandardTextResult(
        action="formatted",
        toc_count=max(len(toc) - 1, 0),
        table_count=table_count,
        merged_cover_label=merged_cover_label,
    )


def _iter_blocks(document: Document) -> Iterable[Block]:
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if text:
                yield ParaBlock(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if any(any(cell for cell in row) for row in rows):
                yield TableBlock(rows)


def _split_standard_document(blocks: list[Block]) -> tuple[list[str], list[str], list[Block]]:
    paragraphs = [block.text for block in blocks if isinstance(block, ParaBlock)]
    compact_texts = [re.sub(r"\s+", "", text) for text in paragraphs]
    toc_start = compact_texts.index("目次") if "目次" in compact_texts else None
    body_start = _find_body_preface_index(paragraphs, toc_start)
    if toc_start is None or body_start is None or body_start <= toc_start:
        raise ValueError("standard text requires a clear 目次 and 前言 structure")

    cover = paragraphs[:toc_start]
    toc = paragraphs[toc_start:body_start]

    body: list[Block] = []
    paragraph_index = -1
    for block in blocks:
        if isinstance(block, ParaBlock):
            paragraph_index += 1
        if paragraph_index >= body_start:
            body.append(block)
    return cover, toc, body


def _find_body_preface_index(paragraphs: list[str], toc_start: int | None) -> int | None:
    if toc_start is None:
        return None
    for index, text in enumerate(paragraphs[toc_start + 1 :], toc_start + 1):
        if re.fullmatch(r"前\s+言", text):
            return index
    compact_matches = [
        index
        for index, text in enumerate(paragraphs[toc_start + 1 :], toc_start + 1)
        if re.sub(r"\s+", "", text) == "前言"
    ]
    return compact_matches[-1] if compact_matches else None


def _merge_cover_standard_label(cover: list[str]) -> tuple[list[str], bool]:
    if len(cover) >= 2 and cover[0] == "中华人民共和国" and "标准" in cover[1]:
        return [cover[0] + cover[1], *cover[2:]], True
    return cover, False


def _add_cover_text(document: Document, index: int, text: str, profile: dict[str, Any]) -> None:
    if index == 0 and "标准" in text:
        _add_text(
            document,
            text,
            profile,
            font_key=SONG_FONT,
            size=16,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent_cm=0,
        )
        _add_empty_line(document, 20)
    elif index in {1, 2}:
        _add_text(
            document,
            text,
            profile,
            font_key=HEADING_FONT,
            size=22,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_pt=32,
            first_indent_cm=0,
        )
    elif "征求意见稿" in text:
        _add_text(
            document,
            text,
            profile,
            font_key=HEADING_FONT,
            size=18,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_pt=30,
            first_indent_cm=0,
        )
        for _ in range(8):
            _add_empty_line(document, 20)
    else:
        _add_text(
            document,
            text,
            profile,
            font_key=SONG_FONT,
            size=16,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent_cm=0,
        )


def _add_toc_text(document: Document, index: int, text: str, profile: dict[str, Any]) -> None:
    compact = re.sub(r"\s+", "", text)
    if index == 0:
        _add_text(
            document,
            text,
            profile,
            font_key=HEADING_FONT,
            size=16,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent_cm=0,
            after=8,
        )
        return

    level = _toc_level(text)
    _add_text(
        document,
        text,
        profile,
        font_key=SONG_FONT,
        size=12,
        bold=(level == 1 and compact != "前言"),
        first_indent_cm=0,
        left_indent_cm=0.74 if level >= 2 else 0,
        line_pt=20,
    )


def _toc_level(text: str) -> int:
    normalized = " ".join(text.split())
    if CLAUSE_RE.match(normalized):
        return 2
    if CHAPTER_RE.match(normalized) or re.sub(r"\s+", "", text) == "前言":
        return 1
    return 2


def _add_body_block(document: Document, block: Block, profile: dict[str, Any]) -> None:
    if isinstance(block, TableBlock):
        _add_table(document, block.rows, profile)
        return

    text = block.text
    compact = re.sub(r"\s+", "", text)
    normalized = " ".join(text.split())
    if compact == "前言":
        _add_text(
            document,
            text,
            profile,
            font_key=HEADING_FONT,
            size=16,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent_cm=0,
            after=10,
        )
    elif CHAPTER_RE.fullmatch(normalized) and not CLAUSE_RE.fullmatch(normalized):
        _add_text(document, text, profile, font_key=HEADING_FONT, size=16, bold=True, first_indent_cm=0, before=8, after=4)
    elif CLAUSE_RE.fullmatch(normalized):
        _add_text(document, text, profile, font_key=HEADING_FONT, size=16, first_indent_cm=0, before=4)
    elif LIST_ITEM_RE.match(text):
        _add_text(document, text, profile, font_key=BODY_FONT, size=16, first_indent_cm=None, left_indent_cm=0.74)
    else:
        _add_text(document, text, profile, font_key=BODY_FONT, size=16, first_indent_cm=0.74)


def _add_text(
    document: Document,
    text: str,
    profile: dict[str, Any],
    *,
    font_key: str = BODY_FONT,
    size: float | None = None,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_pt: float | None = None,
    before: float = 0,
    after: float = 0,
    first_indent_cm: float | None = 0.74,
    left_indent_cm: float | None = None,
):
    paragraph = document.add_paragraph()
    layout = profile.get("layout", {})
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(line_pt if line_pt is not None else float(layout.get("line_spacing_pt", 28)))
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if first_indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_indent_cm)
    if left_indent_cm is not None:
        paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    run = paragraph.add_run(text)
    _set_run(run, profile, font_key, size=size, bold=bold)
    return paragraph


def _add_empty_line(document: Document, line_pt: float = 20) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = Pt(line_pt)


def _add_table(document: Document, rows: list[list[str]], profile: dict[str, Any]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    widths = _column_widths(rows, _content_width_twips(profile))
    _apply_column_widths(table, widths)
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = Pt(20)
            run = paragraph.add_run(text)
            _set_run(run, profile, SONG_FONT, size=10.5, bold=(row_index == 0))
    _add_empty_line(document, 12)


def _column_widths(rows: list[list[str]], total_width_twips: int) -> list[int]:
    column_count = max(len(row) for row in rows)
    if column_count <= 0:
        return []
    header = [cell.strip() for cell in rows[0]]
    if header and header[0] in {"序号", "编号"} and column_count > 1:
        serial_width = int(total_width_twips * 0.12)
        remaining = total_width_twips - serial_width
        content_width = remaining // (column_count - 1)
        return [serial_width, *([content_width] * (column_count - 1))]
    width = total_width_twips // column_count
    return [width] * column_count


def _apply_column_widths(table: Table, widths: list[int]) -> None:
    tbl = table._tbl
    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(0, grid)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index >= len(widths):
                continue
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))


def _set_table_borders(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _add_page_number(section, profile: dict[str, Any]) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)
    _set_run(run, profile, SONG_FONT, size=12)


def _set_page(section, profile: dict[str, Any]) -> None:
    page = profile.get("page", {})
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(float(page.get("top_margin_cm", 2.5)))
    section.bottom_margin = Cm(float(page.get("bottom_margin_cm", 2.5)))
    section.left_margin = Cm(float(page.get("left_margin_cm", 2.8)))
    section.right_margin = Cm(float(page.get("right_margin_cm", 2.6)))


def _set_run(run, profile: dict[str, Any], font_key: str, *, size: float | None = None, bold: bool = False) -> None:
    font = _preferred_font(profile, font_key)
    run.font.name = font
    run.font.size = Pt(size if size is not None else _font_size(profile, font_key, 16))
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)


def _content_width_twips(profile: dict[str, Any]) -> int:
    page = profile.get("page", {})
    page_width_cm = 21.0
    left = float(page.get("left_margin_cm", 2.8))
    right = float(page.get("right_margin_cm", 2.6))
    return int(max(page_width_cm - left - right, 1) * 567)


def _preferred_font(profile: dict[str, Any], key: str) -> str:
    fonts = profile.get("fonts", {}).get(key, {})
    fallbacks = fonts.get("fallbacks") or []
    return fallbacks[0] if fallbacks else "仿宋_GB2312"


def _font_size(profile: dict[str, Any], key: str, default: float) -> float:
    return float(profile.get("fonts", {}).get(key, {}).get("size_pt", default))
