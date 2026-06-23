"""Conservative Word table-of-contents generation."""

from __future__ import annotations

import re
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


LEVEL1_RE = re.compile(r"^([一二三四五六七八九十]+[、，,]|\d+\s+)")
LEVEL2_RE = re.compile(r"^(（[一二三四五六七八九十]+）|\([一二三四五六七八九十]+\))")
LEVEL3_RE = re.compile(r"^\d+(?:[.．]\d+)+\s+|\d+[.．、]")
SENTENCE_RE = re.compile(r"[。；;！？!?]")


@dataclass(frozen=True)
class TocResult:
    heading_count: int
    action: str
    max_level: int = 0


def generate_toc_if_clear(document: Document) -> TocResult:
    """Generate a Word TOC field when numbered headings are clear enough."""

    headings = _detect_headings(document)
    if len(headings) < 2:
        return TocResult(heading_count=len(headings), action="skipped_unclear_headings")

    for paragraph, level in headings:
        _set_outline_level(paragraph, level - 1)
    _insert_toc_at_start(document)
    return TocResult(
        heading_count=len(headings),
        action="generated",
        max_level=max(level for _, level in headings),
    )


def _detect_headings(document: Document) -> list[tuple[object, int]]:
    headings: list[tuple[object, int]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        level = _heading_level(text)
        if level is not None:
            headings.append((paragraph, level))
    return headings


def _heading_level(text: str) -> int | None:
    if not text or len(text) > 34 or SENTENCE_RE.search(text):
        return None
    if LEVEL1_RE.match(text):
        return 1
    if LEVEL3_RE.match(text):
        return 2
    if LEVEL2_RE.match(text):
        return 2
    return None


def _insert_toc_at_start(document: Document) -> None:
    paragraph = document.paragraphs[0].insert_paragraph_before()
    paragraph.add_run("目录")
    toc_paragraph = paragraph.insert_paragraph_before()
    _add_toc_field(toc_paragraph)
    toc_paragraph.add_run().add_break(WD_BREAK.PAGE)


def _add_toc_field(paragraph) -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "右键更新域生成目录"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _set_outline_level(paragraph, level: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    outline_level = paragraph_properties.find(qn("w:outlineLvl"))
    if outline_level is None:
        outline_level = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline_level)
    outline_level.set(qn("w:val"), str(level))
