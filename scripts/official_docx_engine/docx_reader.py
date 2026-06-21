"""Read-only DOCX snapshot extraction."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Optional

from docx import Document

from .models import DocumentSnapshot, ParagraphSnapshot, RunSnapshot


def _pt_value(value: Any) -> Optional[float]:
    if value is None:
        return None
    pt = getattr(value, "pt", None)
    if pt is not None:
        return float(pt)
    return None


def _alignment_name(value: Any) -> Optional[str]:
    if value is None:
        return None
    return getattr(value, "name", str(value))


def _line_spacing_value(value: Any) -> Any:
    if value is None:
        return None
    pt = _pt_value(value)
    if pt is not None:
        return pt
    return value


def read_docx_snapshot(path: str | Path) -> DocumentSnapshot:
    """Read a DOCX file into a neutral snapshot without modifying it."""

    docx_path = Path(path)
    document = Document(str(docx_path))
    paragraphs: list[ParagraphSnapshot] = []

    for index, paragraph in enumerate(document.paragraphs):
        paragraph_format = paragraph.paragraph_format
        style = paragraph.style
        runs = tuple(
            RunSnapshot(
                text=run.text,
                font_name=run.font.name,
                font_size_pt=_pt_value(run.font.size),
                bold=run.bold,
            )
            for run in paragraph.runs
        )
        paragraphs.append(
            ParagraphSnapshot(
                index=index,
                text=paragraph.text,
                style_name=style.name if style is not None else None,
                alignment=_alignment_name(paragraph.alignment),
                left_indent_pt=_pt_value(paragraph_format.left_indent),
                first_line_indent_pt=_pt_value(paragraph_format.first_line_indent),
                line_spacing=_line_spacing_value(paragraph_format.line_spacing),
                line_spacing_pt=_pt_value(paragraph_format.line_spacing),
                runs=runs,
            )
        )

    return DocumentSnapshot(
        path=docx_path,
        paragraphs=tuple(paragraphs),
        table_count=len(document.tables),
    )
