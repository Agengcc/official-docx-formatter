"""Conservative table preservation and internal formatting."""

from __future__ import annotations

import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table


@dataclass(frozen=True)
class TableFormatResult:
    table_count: int
    action: str
    skipped_count: int = 0


def append_and_format_source_tables(
    target_document: Document,
    source_path: str | Path,
    profile: dict[str, Any],
) -> TableFormatResult:
    """Copy source tables into the output document and format cell internals."""

    source_document = Document(str(source_path))
    if not source_document.tables:
        return TableFormatResult(table_count=0, action="no_tables")

    existing_count = len(target_document.tables)
    for table in source_document.tables:
        _append_to_document_body(target_document, copy.deepcopy(table._tbl))

    copied_tables = target_document.tables[existing_count:]
    content_width_twips = _content_width_twips(profile)
    for table in copied_tables:
        _format_table(table, profile, content_width_twips)

    return TableFormatResult(table_count=len(copied_tables), action="formatted")


def copy_and_format_table(
    target_document: Document,
    source_table: Table,
    profile: dict[str, Any],
) -> Table:
    """Copy one source table into the current output position and format it."""

    _append_to_document_body(target_document, copy.deepcopy(source_table._tbl))
    copied_table = target_document.tables[-1]
    _format_table(copied_table, profile, _content_width_twips(profile))
    return copied_table


def _append_to_document_body(target_document: Document, element) -> None:
    body = target_document._body._element
    body.insert_element_before(element, "w:sectPr")


def _format_table(table: Table, profile: dict[str, Any], width_twips: int) -> None:
    table.autofit = not _has_serial_number_column(table)
    _set_table_width(table, width_twips)
    if _has_serial_number_column(table):
        _apply_serial_number_widths(table, width_twips)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_margins(cell, top=80, bottom=80, left=108, right=108)
            for paragraph in cell.paragraphs:
                _format_cell_paragraph(paragraph, profile)


def _format_cell_paragraph(paragraph, profile: dict[str, Any]) -> None:
    body_font = _preferred_font(profile, "body")
    body_size = _font_size(profile, "body", 16)
    layout = profile.get("layout", {})
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(float(layout.get("line_spacing_pt", 28)))
    paragraph.paragraph_format.space_before = Pt(float(layout.get("space_before_pt", 0)))
    paragraph.paragraph_format.space_after = Pt(float(layout.get("space_after_pt", 0)))
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = body_font
        run.font.size = Pt(body_size)
        run.font.bold = False
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), body_font)


def _set_table_width(table: Table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_twips))


def _has_serial_number_column(table: Table) -> bool:
    if not table.rows or not table.columns:
        return False
    first_cell_text = table.cell(0, 0).text.strip()
    return first_cell_text in {"序号", "编号"} and len(table.columns) > 1


def _apply_serial_number_widths(table: Table, total_width_twips: int) -> None:
    column_count = len(table.columns)
    serial_width = int(total_width_twips * 0.12)
    remaining_width = total_width_twips - serial_width
    content_width = remaining_width // (column_count - 1)
    widths = [serial_width, *([content_width] * (column_count - 1))]

    tbl = table._tbl
    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(0, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index >= len(widths):
                continue
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, *, top: int, bottom: int, left: int, right: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        margin = tc_mar.find(qn(f"w:{key}"))
        if margin is None:
            margin = OxmlElement(f"w:{key}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _content_width_twips(profile: dict[str, Any]) -> int:
    page = profile.get("page", {})
    page_width_cm = 21.0
    left = float(page.get("left_margin_cm", 2.8))
    right = float(page.get("right_margin_cm", 2.6))
    return int(max(page_width_cm - left - right, 1) * 567)


def _preferred_font(profile: dict[str, Any], key: str) -> str:
    fonts = profile.get("fonts", {}).get(key, {})
    fallbacks = fonts.get("fallbacks") or []
    return fallbacks[0] if fallbacks else "仿宋"


def _font_size(profile: dict[str, Any], key: str, default: float = 16) -> float:
    return float(profile.get("fonts", {}).get(key, {}).get("size_pt", default))
