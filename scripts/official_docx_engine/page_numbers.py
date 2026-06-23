"""Conservative page-number footer handling for DOCX output."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass(frozen=True)
class FooterInspection:
    existing_page_number: bool = False
    existing_non_page_footer: bool = False
    footer_texts: tuple[str, ...] = ()


@dataclass(frozen=True)
class PageNumberResult:
    existing_page_number: bool
    existing_non_page_footer: bool
    action: str
    preserved_footer_texts: tuple[str, ...] = ()


def inspect_footers(path: str | Path) -> FooterInspection:
    """Inspect source footers without changing the document."""

    document = Document(str(path))
    footer_texts: list[str] = []
    footer_xml_parts: list[str] = []
    for section in document.sections:
        footer = section.footer
        footer_xml_parts.append(footer._element.xml)
        for paragraph in footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                footer_texts.append(text)

    footer_xml = "\n".join(footer_xml_parts).upper()
    has_page_field = "PAGE" in footer_xml and ("FLDCHAR" in footer_xml or "INSTR" in footer_xml)
    has_page_text = any(text.strip().isdigit() for text in footer_texts)
    existing_page_number = has_page_field or has_page_text
    non_page_texts = tuple(text for text in footer_texts if not text.strip().isdigit())
    return FooterInspection(
        existing_page_number=existing_page_number,
        existing_non_page_footer=bool(non_page_texts),
        footer_texts=non_page_texts,
    )


def apply_page_numbers(document: Document, inspection: FooterInspection | None = None) -> PageNumberResult:
    """Insert a PAGE field unless doing so would overwrite non-page footer text."""

    footer_info = inspection or FooterInspection()
    if footer_info.existing_non_page_footer:
        _preserve_footer_texts(document, footer_info.footer_texts)
        return PageNumberResult(
            existing_page_number=footer_info.existing_page_number,
            existing_non_page_footer=True,
            action="skipped_non_page_footer",
            preserved_footer_texts=footer_info.footer_texts,
        )

    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        _add_page_field(paragraph)

    return PageNumberResult(
        existing_page_number=footer_info.existing_page_number,
        existing_non_page_footer=False,
        action="inserted",
        preserved_footer_texts=(),
    )


def _preserve_footer_texts(document: Document, footer_texts: tuple[str, ...]) -> None:
    text = " ".join(footer_texts).strip()
    if not text:
        return
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.text = text


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    display = OxmlElement("w:t")
    display.text = "1"
    run._r.append(display)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
