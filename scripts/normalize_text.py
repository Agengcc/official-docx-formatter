#!/usr/bin/env python3
"""Conservative punctuation and spacing normalization for Chinese official documents."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - environment message
    raise SystemExit("python-docx is required. Install with: pip install python-docx") from exc


PROTECT_PREFIX = "\x02DOCFMT"
PROTECT_SUFFIX = "\x03"

SIMPLE_REPLACEMENTS = {
    "(": "（",
    ")": "）",
    ":": "：",
    ";": "；",
    "?": "？",
    "!": "！",
}

LEFT_DOUBLE_QUOTE = "\u201c"
RIGHT_DOUBLE_QUOTE = "\u201d"


def has_chinese(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def protect_special_patterns(text: str) -> tuple[str, list[tuple[str, str]]]:
    protected: list[tuple[str, str]] = []

    def replace(match: re.Match[str]) -> str:
        placeholder = f"{PROTECT_PREFIX}{len(protected)}{PROTECT_SUFFIX}"
        protected.append((placeholder, match.group(0)))
        return placeholder

    result = text
    patterns = [
        r"(?:https?|ftp)://\S+",
        r"[\w.+-]+@[\w-]+\.[\w.-]+",
        r"[A-Za-z]:\\[^\s，。；：！？、]*",
        r"\b[A-Z]{1,10}(?:/[A-Z])?(?:\s+|-)?\d+(?:[-:]\d+)+\b",
        r"[A-Za-z]+[\s-]?\d+:\d{2,}",
        r"(?<!\d)\d{1,2}:\d{2}(?::\d{2})?(?!\d)",
        r"\b[A-Z]{1,12}[-_/]\d[\w\-_/]*\b",
    ]
    for pattern in patterns:
        result = re.sub(pattern, replace, result)
    return result, protected


def restore_special_patterns(text: str, protected: Iterable[tuple[str, str]]) -> str:
    result = text
    for placeholder, original in protected:
        result = result.replace(placeholder, original)
    return result


def normalize_spaces(text: str, mode: str) -> str:
    if mode == "keep_all" or not text:
        return text
    if mode == "remove_all":
        return text.replace("\u3000", "").replace(" ", "")
    if mode == "keep_en_boundary":
        cn = r"\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff"
        cn_cls = f"[{cn}]"
        en = r"[A-Za-z0-9]"
        text = text.replace("\u3000", " ")
        text = re.sub(f"(?<={cn_cls}) +(?={cn_cls})", "", text)
        text = re.sub(f"(?<={cn_cls}) +(?={en})", " ", text)
        text = re.sub(f"(?<={en}) +(?={cn_cls})", " ", text)
        text = re.sub(f"(?<={cn_cls}) +(?=\\d)", "", text)
        text = re.sub(r"(?<=\d) +(?=[年月日号元万亿％%])", "", text)
        text = re.sub(r"(?<=[第共约]) +(?=\d)", "", text)
        text = re.sub(r"(?<=\d) +(?=[项条个台套件人次])", "", text)
        text = re.sub(r" {2,}", " ", text)
        return text.strip(" ")
    return text


def normalize_quotes(text: str) -> str:
    double_quote_chars = ['"', "\u201c", "\u201d", "\u201e", "\u201f", "\u300c", "\u300d"]

    chars = list(text)
    double_index = 0
    for index, char in enumerate(chars):
        if char in double_quote_chars:
            chars[index] = LEFT_DOUBLE_QUOTE if double_index % 2 == 0 else RIGHT_DOUBLE_QUOTE
            double_index += 1
    return "".join(chars)


def normalize_text(text: str, space_mode: str = "keep_en_boundary") -> str:
    if not text:
        return text

    result, protected = protect_special_patterns(text)

    result = re.sub(r"\.{2,}", "……", result)
    result = re.sub(r"。{2,}", "……", result)
    result = re.sub(r"--+", "——", result)
    result = re.sub(r"(?<!—)—(?!—)", "——", result)

    if has_chinese(result):
        for old, new in SIMPLE_REPLACEMENTS.items():
            result = result.replace(old, new)
    result = re.sub(r"([\u4e00-\u9fff]),", r"\1，", result)
    result = re.sub(r",([\u4e00-\u9fff])", r"，\1", result)
    result = re.sub(r"([\u4e00-\u9fff])\.(\s|$)", r"\1。\2", result)
    result = normalize_quotes(result)
    result = normalize_spaces(result, space_mode)

    return restore_special_patterns(result, protected)


def normalize_docx(input_path: Path, output_path: Path, space_mode: str = "keep_en_boundary") -> int:
    doc = Document(str(input_path))
    changes = 0

    def normalize_paragraphs(paragraphs) -> None:
        nonlocal changes
        for paragraph in paragraphs:
            original = paragraph.text
            normalized = normalize_text(original, space_mode=space_mode)
            if normalized != original:
                paragraph.text = normalized
                changes += 1

    normalize_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_paragraphs(cell.paragraphs)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return changes


def main() -> int:
    parser = argparse.ArgumentParser(description="Normalize punctuation and spacing in a DOCX file")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("-o", "--output", required=True, help="Output .docx file")
    parser.add_argument(
        "--space-mode",
        choices=("keep_en_boundary", "remove_all", "keep_all"),
        default="keep_en_boundary",
        help="How to handle spaces in Chinese text.",
    )
    args = parser.parse_args()

    changes = normalize_docx(Path(args.input), Path(args.output), space_mode=args.space_mode)
    print(f"saved={args.output}")
    print(f"normalized_paragraphs={changes}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
