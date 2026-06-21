#!/usr/bin/env python3
"""Format Chinese official DOCX documents using JSON profiles."""

from __future__ import annotations

import argparse
import copy
import json
import re
from pathlib import Path
from typing import Any, Dict, Iterable, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - environment message
    raise SystemExit("python-docx is required. Install with: pip install python-docx") from exc

SKILL_DIR = Path(__file__).resolve().parents[1]
PROFILES_DIR = SKILL_DIR / "profiles"
DOCUMENT_TYPES_FILE = SKILL_DIR / "references" / "document_types.json"
BASE_PROFILE = "standard-party-government"

CN_DATE_RE = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日$")


def _engine_imports():
    from official_docx_engine.diagnostics import diagnose_snapshot
    from official_docx_engine.docx_reader import read_docx_snapshot
    from official_docx_engine.format_plan import build_format_plan
    from official_docx_engine.reporting import write_report_json
    from official_docx_engine.structure import analyze_structure

    return read_docx_snapshot, analyze_structure, diagnose_snapshot, build_format_plan, write_report_json


def deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
    result = copy.deepcopy(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = copy.deepcopy(value)
    return result


def load_profile(profile_id: str) -> Dict[str, Any]:
    path = PROFILES_DIR / f"{profile_id}.json"
    if not path.exists():
        raise SystemExit(f"profile not found: {profile_id}")
    profile = json.loads(path.read_text(encoding="utf-8"))
    parent_id = profile.get("inherits")
    if parent_id:
        parent = load_profile(parent_id)
        profile = deep_merge(parent, profile)
    return profile


def load_document_types() -> Dict[str, Any]:
    return json.loads(DOCUMENT_TYPES_FILE.read_text(encoding="utf-8"))


def classify_lines_for_type(lines: list[str]) -> Dict[str, Any]:
    import sys

    scripts_dir = str(Path(__file__).resolve().parent)
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from classify_document import classify_lines

    return classify_lines(lines)


def normalize_line(text: str, enabled: bool, space_mode: str) -> str:
    if not enabled:
        return text
    import sys

    scripts_dir = str(Path(__file__).resolve().parent)
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from normalize_text import normalize_text

    return normalize_text(text, space_mode=space_mode)


def build_skeleton_body(doc_type: str) -> list[str]:
    document_types = load_document_types()
    if doc_type not in document_types:
        supported = "、".join(document_types)
        raise SystemExit(f"unsupported doc type: {doc_type}; supported: {supported}")

    info = document_types[doc_type]
    body: list[str] = []
    for index, section in enumerate(info.get("sections", []), 1):
        body.append(f"{_cn_number(index)}、{section}")
        body.append(f"请补充{section}。")
    ending = info.get("ending")
    if ending:
        body.append(ending)
    return body


def _cn_number(value: int) -> str:
    numbers = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 0 <= value <= 10:
        return numbers[value]
    return str(value)


def preferred_font(profile: Dict[str, Any], key: str) -> str:
    fonts = profile.get("fonts", {}).get(key, {})
    fallbacks = fonts.get("fallbacks") or []
    return fallbacks[0] if fallbacks else "仿宋"


def font_size(profile: Dict[str, Any], key: str, default: float = 16) -> float:
    return float(profile.get("fonts", {}).get(key, {}).get("size_pt", default))


def font_bold(profile: Dict[str, Any], key: str) -> bool:
    return bool(profile.get("fonts", {}).get(key, {}).get("bold", False))


def add_run(paragraph, text: str, font: str, size_pt: float, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_common_paragraph_format(paragraph, profile: Dict[str, Any], first_line: bool = True) -> None:
    layout = profile.get("layout", {})
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(float(layout.get("line_spacing_pt", 28)))
    paragraph.paragraph_format.space_before = Pt(float(layout.get("space_before_pt", 0)))
    paragraph.paragraph_format.space_after = Pt(float(layout.get("space_after_pt", 0)))
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(32)
    else:
        paragraph.paragraph_format.first_line_indent = Pt(0)


def setup_page(doc: Document, profile: Dict[str, Any]) -> None:
    page = profile.get("page", {})
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(float(page.get("top_margin_cm", 3.7)))
    section.bottom_margin = Cm(float(page.get("bottom_margin_cm", 3.5)))
    section.left_margin = Cm(float(page.get("left_margin_cm", 2.8)))
    section.right_margin = Cm(float(page.get("right_margin_cm", 2.6)))


def looks_like_inline_body(text: str) -> bool:
    stripped = text.strip()
    if re.search(r"[。；;：:]", stripped):
        return True
    return len(stripped) > 34


def hierarchy_key(text: str) -> Optional[str]:
    if re.match(r"^[一二三四五六七八九十][、，,]", text):
        return "level1"
    if re.match(r"^(（[一二三四五六七八九十]）|\([一二三四五六七八九十]\))", text):
        return None if looks_like_inline_body(text) else "level2"
    if re.match(r"^\d+[\.．]", text):
        return None if looks_like_inline_body(text) else "body"
    if re.match(r"^(（\d+）|\(\d+\))", text):
        return None if looks_like_inline_body(text) else "body"
    return None


def split_source_document(input_path: Path, normalize: bool = True, space_mode: str = "keep_en_boundary") -> Tuple[str, str, list[str], Optional[str], Optional[str]]:
    source = Document(str(input_path))
    paragraphs = [
        normalize_line(p.text.strip(), normalize, space_mode)
        for p in source.paragraphs
        if p.text.strip()
    ]
    title = paragraphs[0] if paragraphs else ""
    recipient = ""
    body_start = 1

    for index, text in enumerate(paragraphs[1:], 1):
        if text.endswith(("：", ":")) and len(text) <= 40:
            recipient = text.rstrip("：:").strip()
            body_start = index + 1
            break

    body = paragraphs[body_start:]
    issuer = None
    date_text = None
    if len(body) >= 2 and CN_DATE_RE.match(body[-1]):
        issuer = body[-2]
        date_text = body[-1]
        body = body[:-2]
    return title, recipient, body, issuer, date_text


def add_title(doc: Document, title: str, profile: Dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    add_run(paragraph, title, preferred_font(profile, "title"), font_size(profile, "title", 22), font_bold(profile, "title"))


def add_recipient(doc: Document, recipient: str, profile: Dict[str, Any]) -> None:
    if not recipient:
        return
    paragraph = doc.add_paragraph()
    set_common_paragraph_format(paragraph, profile, first_line=False)
    add_run(paragraph, f"{recipient}：", preferred_font(profile, "body"), font_size(profile, "body", 16), False)


def add_body_paragraph(doc: Document, text: str, profile: Dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    key = hierarchy_key(text) or "body"
    first_line = key not in {"level1"}
    set_common_paragraph_format(paragraph, profile, first_line=first_line)
    add_run(paragraph, text, preferred_font(profile, key), font_size(profile, key, 16), font_bold(profile, key))


def add_footer(doc: Document, issuer: Optional[str], date_text: Optional[str], profile: Dict[str, Any]) -> None:
    if not issuer and not date_text:
        return
    signature = profile.get("signature", {})
    blank_lines_before = int(signature.get("blank_lines_before", 1))
    right_empty_chars = float(signature.get("right_empty_chars", 2))
    body_font = preferred_font(profile, "body")
    body_size = font_size(profile, "body", 16)
    for _ in range(max(blank_lines_before, 0)):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = Pt(float(profile.get("layout", {}).get("line_spacing_pt", 28)))
        add_run(paragraph, "", body_font, body_size, False)
    for value in [issuer, date_text]:
        if not value:
            continue
        paragraph = doc.add_paragraph()
        set_common_paragraph_format(paragraph, profile, first_line=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.right_indent = Pt(body_size * right_empty_chars)
        add_run(paragraph, value, body_font, body_size, False)


def build_document(title: str, recipient: str, body: Iterable[str], issuer: Optional[str], date_text: Optional[str], profile: Dict[str, Any], normalize: bool = False, space_mode: str = "keep_en_boundary") -> Document:
    doc = Document()
    setup_page(doc, profile)
    title = normalize_line(title, normalize, space_mode)
    recipient = normalize_line(recipient, normalize, space_mode)
    issuer = normalize_line(issuer, normalize, space_mode) if issuer else issuer
    date_text = normalize_line(date_text, normalize, space_mode) if date_text else date_text
    if title:
        add_title(doc, title, profile)
    add_recipient(doc, recipient, profile)
    for text in body:
        text = normalize_line(text.strip(), normalize, space_mode)
        if text:
            add_body_paragraph(doc, text, profile)
    add_footer(doc, issuer, date_text, profile)
    return doc


def smoke_check(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return f"paragraphs={len(paragraphs)}"


def _report_path(output_path: Path, report_arg: Optional[str]) -> Optional[Path]:
    if report_arg in {None, ""}:
        return output_path.with_suffix(".report.json")
    return Path(report_arg)


def main() -> int:
    parser = argparse.ArgumentParser(description="Format Chinese official DOCX documents")
    parser.add_argument("input", nargs="?", help="Input .docx path. Omit when using --text-file.")
    parser.add_argument("-o", "--output", required=True, help="Output .docx path")
    parser.add_argument("--profile", default=BASE_PROFILE, help="Profile id from profiles/*.json")
    parser.add_argument("--title", help="Title when creating from text")
    parser.add_argument("--recipient", default="", help="Recipient when creating from text")
    parser.add_argument("--issuer", help="Issuer override")
    parser.add_argument("--date", dest="date_text", help="Date override, e.g. 2026年6月19日")
    parser.add_argument("--text-file", help="Create a formatted .docx from a plain text file")
    parser.add_argument("--doc-type", help="Document type, e.g. 请示, 报告, 通知, 函, 纪要")
    parser.add_argument("--create-skeleton", action="store_true", help="Create a document skeleton from --doc-type")
    parser.add_argument("--assume-detected-type", action="store_true", help="Continue when classifier is confident; otherwise stop and ask user")
    parser.add_argument(
        "--report",
        nargs="?",
        const="",
        default=None,
        help="Write a JSON formatting report for DOCX input. Defaults to output_path.with_suffix('.report.json').",
    )
    parser.add_argument("--no-normalize-text", action="store_true", help="Do not normalize punctuation or spacing")
    parser.add_argument(
        "--space-mode",
        choices=("keep_en_boundary", "remove_all", "keep_all"),
        default="keep_en_boundary",
        help="How to handle spaces when normalizing text.",
    )
    args = parser.parse_args()

    profile = load_profile(args.profile)
    output_path = Path(args.output)
    report_path = _report_path(output_path, args.report)
    normalize = not args.no_normalize_text
    report_data = None

    if args.create_skeleton:
        if not args.doc_type:
            raise SystemExit("--doc-type is required with --create-skeleton")
        title = args.title or f"关于××事项的{args.doc_type}"
        body = build_skeleton_body(args.doc_type)
        doc = build_document(title, args.recipient, body, args.issuer, args.date_text, profile, normalize=normalize, space_mode=args.space_mode)
    elif args.text_file:
        text = Path(args.text_file).read_text(encoding="utf-8")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        title = args.title or (lines[0] if lines else "")
        body = lines[1:] if title and lines and title == lines[0] else lines
        doc = build_document(title, args.recipient, body, args.issuer, args.date_text, profile, normalize=normalize, space_mode=args.space_mode)
    else:
        if not args.input:
            raise SystemExit("input .docx path is required unless --text-file is used")
        input_path = Path(args.input)
        read_docx_snapshot, analyze_structure, diagnose_snapshot, build_format_plan, write_report_json = _engine_imports()
        snapshot = read_docx_snapshot(input_path)
        source_lines = [paragraph.text.strip() for paragraph in snapshot.non_empty_paragraphs]
        if not args.doc_type:
            classification = classify_lines_for_type(source_lines)
            top = classification["top"]
            print(f"detected_doc_type={top.get('doc_type', '未知')} confidence={top.get('confidence', 0)}")
            if classification["ask_user"] and not args.assume_detected_type:
                print(classification["question"])
                return 2
            args.doc_type = top.get("doc_type")

        structure = analyze_structure(snapshot)
        issues = diagnose_snapshot(snapshot, structure)
        plan = build_format_plan(
            snapshot,
            structure,
            profile_id=profile.get("profile_id", args.profile),
            doc_type=args.doc_type or "未知",
            normalize_text=normalize,
        )
        title, recipient, body, detected_issuer, detected_date = split_source_document(input_path, normalize=normalize, space_mode=args.space_mode)
        doc = build_document(
            args.title or title,
            args.recipient or recipient,
            body,
            args.issuer or detected_issuer,
            args.date_text or detected_date,
            profile,
            normalize=False,
            space_mode=args.space_mode,
        )
        if report_path is not None:
            report_data = {
                "write_report_json": write_report_json,
                "input_path": input_path,
                "profile_id": profile.get("profile_id", args.profile),
                "doc_type": args.doc_type or "未知",
                "structure": structure,
                "issues": issues,
                "operations": plan.operations,
            }

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    if report_path is not None and report_data is not None:
        report_data["write_report_json"](
            report_path,
            input_path=report_data["input_path"],
            output_path=output_path,
            profile_id=report_data["profile_id"],
            doc_type=report_data["doc_type"],
            structure=report_data["structure"],
            issues=report_data["issues"],
            operations=report_data["operations"],
        )
        print(f"report={report_path}")
    print(f"saved={output_path}")
    print(f"profile={profile.get('profile_id', args.profile)}")
    print(f"text_normalization={'off' if args.no_normalize_text else args.space_mode}")
    print(smoke_check(output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
